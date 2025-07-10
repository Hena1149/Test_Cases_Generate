# Chargement du modèle spaCy avec gestion d'erreur améliorée
try:
    import spacy
    try:
        nlp = spacy.load("fr_core_news_sm")
    except OSError:
        import subprocess
        import sys
        subprocess.run([sys.executable, "-m", "spacy", "download", "fr_core_news_sm"], check=True)
        nlp = spacy.load("fr_core_news_sm")
except Exception as e:
    st.error(f"Erreur de chargement du modèle spaCy: {str(e)}")
    nlp = None  # Gestion élégante si le modèle ne peut pas être chargé

# Autres imports
from utils.text_processing import is_similar, remove_duplicates, generate_wordcloud, clean_text
from utils.file_utils import process_uploaded_file, export_to_excel, export_test_cases_to_excel
from utils.openai_utils import split_text, generate_rules, generate_checkpoints, generate_test_cases
from collections import Counter
import streamlit as st
import os
import tempfile
import io
from docx import Document
import re
from difflib import SequenceMatcher
from datetime import datetime
import PyPDF2
from tqdm import tqdm

def is_similar(text1: str, text2: str, threshold: float = 0.85) -> bool:
    """Détermine si deux textes sont similaires."""
    return SequenceMatcher(None, text1.lower(), text2.lower()).ratio() >= threshold

# Configuration de la page
st.set_page_config(
    page_title="Génération automatique des cas de tests",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Style CSS personnalisé
st.markdown("""
    <style>
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        font-family: 'Arial', sans-serif;
    }
    .main {
        background-color: #f8f9fa;
    }
    .stButton>button {
        background-color: #4CAF50;
        color: white;
        border-radius: 5px;
        padding: 0.5rem 1rem;
    }
    .stFileUploader>div>div>button {
        background-color: #2196F3;
        color: white;
    }
    .sidebar .sidebar-content {
        background-color: #e3f2fd;
    }
    h1 {
        color: #2c3e50;
    }
    .progress-bar {
        margin-bottom: 1rem;
    }
    .progress-text {
        font-size: 0.8rem;
        color: #666;
        margin-top: -10px;
        margin-bottom: 10px;
    }
    </style>
    """, unsafe_allow_html=True)

def show_progress(current, total, message):
    """Affiche une barre de progression améliorée avec pourcentage."""
    progress = current / total
    percent = int(progress * 100)
    progress_bar = st.progress(progress)
    progress_text = "{}% - {} ({}/{})".format(percent, message, current, total)
    progress_bar.progress(progress, text=progress_text)
    if current == total:
        progress_bar.empty()
        st.toast("Tâche terminée : {}".format(message), icon="✅")

def main():
    st.title("Génération automatique des cas de tests")
    st.markdown("""
    Chargez votre cahier de charge (PDF ou Word) pour en extraire :
    - Les règles de gestion
    - Les points de contrôle
    - Les cas de test
    """)

    # Initialisation des variables de session
    session_defaults = {
        'text': "",
        'rules': [],
        'checkpoints': [],
        'test_cases': [],
        'existing_checkpoints': []
    }
    
    for key, value in session_defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

    # Sidebar pour les paramètres
    with st.sidebar:
        st.header("Paramètres")
        st.session_state.openai_key = st.text_input("Clé API OpenAI", type="password")
        st.session_state.openai_endpoint = st.text_input("Endpoint Azure OpenAI", "https://chat-genai.openai.azure.com/")
        st.session_state.model_name = st.selectbox("Modèle", ["gpt-4o", "gpt-35-turbo"])
        st.divider()
        st.info("Configurez votre clé API et endpoint avant de commencer.")

    # Onglets principaux
    tab1, tab2, tab3, tab4 = st.tabs(["Upload", "Analyse", "Points de contrôle", "Cas de test"])

    with tab1:
        st.header("Chargement du document")
        uploaded_file = st.file_uploader("Téléversez votre cahier des charges", type=["pdf", "docx", "txt"])
        
        if uploaded_file is not None:
            with st.spinner("Extraction du texte en cours..."):
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_path = tmp_file.name
                
                st.session_state.text = process_uploaded_file(tmp_path)
                os.unlink(tmp_path)
            
            st.success("Texte extrait avec succès !")
            with st.expander("Aperçu du texte extrait"):
                st.text(st.session_state.text[:2000] + "...")

    with tab2:
        st.header("Analyse Textuelle")
        
        if not st.session_state.text:
            st.warning("Veuillez d'abord charger un document dans l'onglet Upload.")
            st.stop()
        
        # Analyse textuelle de base
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Nuage de mots clés")
            with st.spinner("Génération du wordcloud..."):
                fig = generate_wordcloud(st.session_state.text)
                st.pyplot(fig)
        
        with col2:
            st.subheader("Mots les plus fréquents")
            tokens = clean_text(st.session_state.text)
            freq_dist = Counter(tokens)
            top_words = freq_dist.most_common(10)
            
            for word, freq in top_words:
                st.markdown("- **{}**: {} occurrences".format(word, freq))
            
            st.download_button(
                label="Télécharger l'analyse",
                data="\n".join(["{}: {}".format(w, f) for w, f in top_words]),
                file_name="frequence_mots.txt",
                key="download_word_freq"
            )

        # Génération des règles
        st.divider()
        st.subheader("Génération des règles de gestion")
        
        if st.button("Générer les règles", type="primary", key="gen_rules_btn"):
            with st.spinner("Analyse en cours avec IA..."):
                try:
                    chunks = [st.session_state.text[i:i+4000] for i in range(0, len(st.session_state.text), 4000)]
                    all_rules = []
                    
                    progress_bar = st.progress(0)
                    for i, chunk in enumerate(chunks, 1):
                        progress = i / len(chunks)
                        percent = int(progress * 100)
                        progress_bar.progress(progress, text="{}% - Traitement chunk {}/{}".format(percent, i, len(chunks)))
                        rules = generate_rules(
                            chunk,
                            st.session_state.openai_key,
                            st.session_state.openai_endpoint,
                            st.session_state.model_name
                        )
                        all_rules.extend(rules)
                    
                    st.session_state.rules = [rule.strip() for rule in all_rules if rule.strip()]
                    progress_bar.empty()
                    st.success("{} règles générées avec succès !".format(len(st.session_state.rules)))
                except Exception as e:
                    st.error("Erreur lors de la génération : {}".format(str(e)))

        # Affichage et export des règles
        if st.session_state.rules:
            st.divider()
            
            with st.expander("Aperçu des {} règles".format(len(st.session_state.rules)), expanded=True):
                show_rules = st.slider(
                    "Nombre de règles à afficher",
                    5, min(50, len(st.session_state.rules)), 10,
                    key="rules_slider"
                )
                
                for i, rule in enumerate(st.session_state.rules[:show_rules], 1):
                    st.markdown("**{}.** {}".format(i, rule))
                
                if len(st.session_state.rules) > show_rules:
                    st.info("Affichage de {}/{} règles".format(show_rules, len(st.session_state.rules)))
            
            # Export multi-format
            st.subheader("Exporter les règles")
            export_format = st.radio("Format d'export", ["Word (.docx)", "Texte (.txt)", "Excel (.xlsx)"], 
                                   horizontal=True, key="rules_export_format")
            
            if export_format == "Word (.docx)":
                try:
                    doc = Document()
                    doc.add_heading('Règles de Gestion', 0)
                    
                    for rule in st.session_state.rules:
                        doc.add_paragraph(rule, style='ListBullet')
                    
                    docx_bytes = io.BytesIO()
                    doc.save(docx_bytes)
                    docx_bytes.seek(0)
                    
                    st.download_button(
                        label="Télécharger (.docx)",
                        data=docx_bytes.getvalue(),
                        file_name="regles_gestion.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_rules_docx"
                    )
                except Exception as e:
                    st.error("Erreur DOCX : {}".format(str(e)))
            
            elif export_format == "Texte (.txt)":
                try:
                    txt_content = "RÈGLES DE GESTION\n\n" + \
                                "\n".join("{}. {}".format(i+1, r) for i, r in enumerate(st.session_state.rules))
                    
                    st.download_button(
                        label="Télécharger (.txt)",
                        data=txt_content,
                        file_name="regles_gestion.txt",
                        mime="text/plain",
                        key="download_rules_txt"
                    )
                except Exception as e:
                    st.error("Erreur TXT : {}".format(str(e)))
            
            elif export_format == "Excel (.xlsx)":
                try:
                    excel_data = export_to_excel(st.session_state.rules, "Regles_gestion")
                    st.download_button(
                        label="📊 Télécharger (.xlsx)",
                        data=excel_data,
                        file_name="regles_gestion.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="download_rules_excel"
                    )
                except Exception as e:
                    st.error("Erreur Excel : {}".format(str(e)))

    with tab3:
        st.header("Points de Contrôle", divider="blue")

        if not st.session_state.text:
            st.warning("Veuillez d'abord charger un document dans l'onglet Upload.")
            st.stop()
        
        # Génération directe à partir du texte
        st.subheader("Génération directe à partir du texte")
        if st.button("Générer les points de contrôle à partir du texte", 
                    type="primary",
                    key="gen_cp_from_text"):
            with st.spinner("Analyse du texte pour générer les points de contrôle..."):
                try:
                    progress_bar = st.progress(0, text="0% - Préparation...")
                    chunks = [st.session_state.text[i:i+4000] for i in range(0, len(st.session_state.text), 4000)]
                    all_points = []
                    
                    for i, chunk in enumerate(chunks):
                        percent = int((i + 1) / len(chunks) * 100)
                        progress_bar.progress(percent / 100, text="{}% - Traitement du chunk {}/{}".format(percent, i+1, len(chunks)))
                        
                        points = generate_checkpoints(
                            [chunk],
                            st.session_state.openai_key,
                            st.session_state.openai_endpoint,
                            st.session_state.model_name
                        )
                        all_points.extend(points)
                    
                    final_points = remove_duplicates(all_points, st.session_state.existing_checkpoints)
                    st.session_state.checkpoints = st.session_state.existing_checkpoints + final_points
                    st.success("{} points de contrôle générés directement à partir du texte !".format(len(final_points)))
                except Exception as e:
                    st.error("Échec de la génération : {}".format(str(e)))
                finally:
                    progress_bar.empty()

        # Génération à partir des règles
        st.divider()
        st.subheader("Génération à partir des règles de gestion")
        
        if not st.session_state.rules:
            st.warning("Aucune règle de gestion disponible. Vous pouvez en générer dans l'onglet 'Analyse'.")
        else:
            if st.button("Générer les points de contrôle à partir des règles", 
                        type="primary",
                        key="gen_cp_from_rules"):
                with st.spinner("Transformation des règles en points vérifiables..."):
                    try:
                        progress_bar = st.progress(0, text="0% - Préparation...")
                        total_rules = len(st.session_state.rules)
                        new_points = []
                        batch_size = 5
                        
                        for i in range(0, total_rules, batch_size):
                            processed = min(i + batch_size, total_rules)
                            percent = int(processed / total_rules * 100)
                            progress_bar.progress(percent / 100, text="{}% - Traitement des règles {}/{}".format(percent, processed, total_rules))
                            
                            batch = st.session_state.rules[i:i + batch_size]
                            points = generate_checkpoints(
                                batch,
                                st.session_state.openai_key,
                                st.session_state.openai_endpoint,
                                st.session_state.model_name
                            )
                            new_points.extend(points)
                        
                        final_points = remove_duplicates(new_points, st.session_state.existing_checkpoints)
                        st.session_state.checkpoints = st.session_state.existing_checkpoints + final_points
                        st.success("{} points de contrôle générés à partir des règles !".format(len(final_points)))
                    except Exception as e:
                        st.error("Échec de la génération : {}".format(str(e)))
                    finally:
                        progress_bar.empty()

        # Import des points existants
        st.subheader("Importer des points existants (facultatif)")
        existing_cp_file = st.file_uploader(
            "Téléverser un fichier de points existants",
            type=["pdf", "docx", "txt"],
            key="existing_cp_upload",
            label_visibility="collapsed"
        )
    
        if existing_cp_file:
            with st.spinner("Analyse du fichier en cours..."):
                try:
                    if existing_cp_file.type == "text/plain":
                        content = existing_cp_file.getvalue().decode("utf-8")
                    elif existing_cp_file.type == "application/pdf":
                        reader = PyPDF2.PdfReader(existing_cp_file)
                        content = "\n".join([page.extract_text() for page in reader.pages])
                    elif existing_cp_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        doc = Document(existing_cp_file)
                        content = "\n".join([para.text for para in doc.paragraphs])
                
                    pattern = r"^(Vérifier|S['']?assurer|Verifier|►|•|\d+[.)])\s+"
                    points = []
                    for line in content.split('\n'):
                        line = line.strip()
                        if re.match(pattern, line, re.IGNORECASE):
                            clean_point = re.sub(pattern, "", line, flags=re.IGNORECASE)
                            if clean_point:
                                points.append(clean_point)
                
                    if points:
                        st.session_state.existing_checkpoints = points
                        st.success("✅ {} points valides détectés".format(len(points)))
                    else:
                        st.warning("Aucun point de contrôle valide détecté dans le fichier")
                except Exception as e:
                    st.error("Erreur lors de l'extraction : {}".format(str(e)))

        # Visualisation des points
        if st.session_state.checkpoints:
            st.subheader("Visualisation des points")
            
            with st.expander("Filtres", expanded=False):
                search_term = st.text_input("Recherche textuelle", key="cp_search")
                col_sort, col_filter = st.columns(2)
                with col_sort:
                    sort_order = st.selectbox("Trier par", ["Ordre original", "Ordre alphabétique"], key="sort_order_cp")
                with col_filter:
                    filter_type = st.selectbox("Filtrer par", ["Tous", "Existants uniquement", "Nouveaux uniquement"], key="filter_type_cp")
            
            filtered_points = st.session_state.checkpoints.copy()
            
            if search_term:
                filtered_points = [p for p in filtered_points if search_term.lower() in p.lower()]
            
            if filter_type == "Existants uniquement":
                filtered_points = [p for p in filtered_points if p in st.session_state.existing_checkpoints]
            elif filter_type == "Nouveaux uniquement":
                filtered_points = [p for p in filtered_points if p not in st.session_state.existing_checkpoints]
            
            if sort_order == "Ordre alphabétique":
                filtered_points.sort(key=lambda x: x.lower())
            
            items_per_page = st.slider("Points par page", 5, 50, 10, key="cp_per_page")
            total_pages = max(1, (len(filtered_points) + items_per_page - 1) // items_per_page)
            page = st.number_input("Page", 1, total_pages, 1, key="cp_page")
            start_idx = (page - 1) * items_per_page
            end_idx = start_idx + items_per_page
            
            for i, point in enumerate(filtered_points[start_idx:end_idx], start=start_idx+1):
                is_existing = point in st.session_state.existing_checkpoints
                
                st.markdown("""
                <div style='
                    padding:10px;
                    margin:5px 0;
                    border-left:4px solid {};
                    background:#f8f9fa;
                    border-radius:5px;
                    box-shadow:0 1px 2px rgba(0,0,0,0.1)'
                >
                    <div style='font-weight:bold; margin-bottom:3px'>
                        Point {} {}
                    </div>
                    <div>{}</div>
                </div>
                """.format("#4CAF50" if not is_existing else "#2196F3",
                          i,
                          "(nouveau)" if not is_existing else "(existant)",
                          point), unsafe_allow_html=True)
            
            st.caption("Page {}/{} • {} points filtrés • {} points au total".format(
                page, total_pages, len(filtered_points), len(st.session_state.checkpoints)))

            # Export des points
            st.subheader("Exporter les points")
            export_format = st.radio("Format d'export", ["Word (.docx)", "Texte (.txt)", "Excel (.xlsx)"], 
                                   horizontal=True, key="cp_export_format")
            
            if export_format == "Word (.docx)":
                try:
                    doc = Document()
                    doc.add_heading('Points de Contrôle', level=1)
                    
                    if st.session_state.existing_checkpoints:
                        doc.add_heading('Points Existants', level=2)
                        for point in st.session_state.existing_checkpoints:
                            doc.add_paragraph(point, style='ListBullet')
                    
                    existing_set = set(st.session_state.existing_checkpoints)
                    new_points = [p for p in st.session_state.checkpoints if p not in existing_set]
                    
                    if new_points:
                        doc.add_heading('Nouveaux Points', level=2)
                        for point in new_points:
                            doc.add_paragraph(point, style='ListBullet')
                    
                    doc_bytes = io.BytesIO()
                    doc.save(doc_bytes)
                    doc_bytes.seek(0)
                    
                    st.download_button(
                        label="Télécharger le DOCX",
                        data=doc_bytes.getvalue(),
                        file_name="points_controle.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_cp_docx"
                    )
                except Exception as e:
                    st.error("Erreur DOCX : {}".format(str(e)))
            
            elif export_format == "Texte (.txt)":
                try:
                    content = "POINTS DE CONTRÔLE\n\n"
                    content += "Généré le {}\n\n".format(datetime.now().strftime('%d/%m/%Y à %H:%M'))
                    
                    if st.session_state.existing_checkpoints:
                        content += "=== POINTS EXISTANTS ===\n"
                        content += "\n".join("• {}".format(p) for p in st.session_state.existing_checkpoints) + "\n\n"
                    
                    existing_set = set(st.session_state.existing_checkpoints)
                    new_points = [p for p in st.session_state.checkpoints if p not in existing_set]
                    
                    if new_points:
                        content += "=== NOUVEAUX POINTS ===\n"
                        content += "\n".join("• {}".format(p) for p in new_points)
                    
                    st.download_button(
                        label="Télécharger le TXT",
                        data=content,
                        file_name="points_controle.txt",
                        mime="text/plain",
                        key="download_cp_txt"
                    )
                except Exception as e:
                    st.error("Erreur TXT : {}".format(str(e)))
            
            elif export_format == "Excel (.xlsx)":
                try:
                    excel_data = export_to_excel(st.session_state.checkpoints, "Points_de_controle")
                    st.download_button(
                        label="📊 Télécharger (.xlsx)",
                        data=excel_data,
                        file_name="points_controle.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="download_cp_excel"
                    )
                except Exception as e:
                    st.error("Erreur Excel : {}".format(str(e)))

    with tab4:
        st.header("Cas de Test")
        
        if not st.session_state.checkpoints:
            st.warning("Veuillez d'abord générer des points de contrôle dans l'onglet précédent.")
        else:
            if st.button("Générer les cas de test", 
                        type="primary",
                        key="gen_tests_from_points"):
                with st.spinner("Création des cas de test..."):
                    try:
                        progress_bar = st.progress(0, text="0% - Préparation...")
                        st.session_state.test_cases = []
                        total = len(st.session_state.checkpoints)
                        
                        test_cases = []
                        for i, checkpoint in enumerate(st.session_state.checkpoints, 1):
                            percent = int(i / total * 100)
                            progress_bar.progress(percent / 100, text="{}% - Génération du cas {}/{}".format(percent, i, total))
                            
                            test_case = generate_test_cases(
                                [checkpoint],
                                st.session_state.openai_key,
                                st.session_state.openai_endpoint,
                                st.session_state.model_name
                            )
                            test_cases.extend(test_case)
                        
                        st.session_state.test_cases = test_cases
                        progress_bar.empty()
                        st.success("{} cas de test générés !".format(len(st.session_state.test_cases)))
                    except Exception as e:
                        st.error("Erreur de génération : {}".format(str(e)))

            if st.session_state.test_cases:
                selected_case = st.selectbox(
                    "Sélectionnez un cas à visualiser",
                    range(len(st.session_state.test_cases)),
                    format_func=lambda x: "Cas de test #{}".format(x+1),
                    key="select_test_case"
                )
                
                st.markdown(st.session_state.test_cases[selected_case])
                
                # Export
                st.subheader("Exporter les cas de test")
                export_format = st.radio("Format d'export", ["Word (.docx)", "Texte (.txt)", "Excel (.xlsx)"], 
                                       horizontal=True, key="test_export_format")
                
                if export_format == "Word (.docx)":
                    try:
                        doc = Document()
                        doc.add_heading('Cas de Test', level=1)
                        
                        for i, test_case in enumerate(st.session_state.test_cases, 1):
                            cleaned_text = re.sub(r'#+\s*', '', test_case)
                            cleaned_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cleaned_text)
                            doc.add_paragraph("Cas de test {}".format(i), style='Heading2')
                            doc.add_paragraph(cleaned_text)
                        
                        docx_bytes = io.BytesIO()
                        doc.save(docx_bytes)
                        docx_bytes.seek(0)
                        
                        st.download_button(
                            label="Télécharger (.docx)",
                            data=docx_bytes.getvalue(),
                            file_name="cas_de_test.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_tests_docx"
                        )
                    except Exception as e:
                        st.error("Erreur DOCX : {}".format(str(e)))
                
                elif export_format == "Texte (.txt)":
                    try:
                        txt_content = "CAS DE TEST\n\n"
                        txt_content += "Généré le {}\n\n".format(datetime.now().strftime('%d/%m/%Y à %H:%M'))
                        txt_content += "\n\n".join(
                            "=== CAS DE TEST {} ===\n{}".format(i+1, re.sub(r'#+\s*|\*\*', '', case))
                            for i, case in enumerate(st.session_state.test_cases)
                        )
                        
                        st.download_button(
                            label="📄 Télécharger (.txt)",
                            data=txt_content,
                            file_name="cas_de_test.txt",
                            mime="text/plain",
                            key="download_tests_txt"
                        )
                    except Exception as e:
                        st.error("Erreur TXT : {}".format(str(e)))
                
                elif export_format == "Excel (.xlsx)":
                    try:
                        excel_data = export_test_cases_to_excel(st.session_state.test_cases)
                        st.download_button(
                            label="📊 Télécharger (.xlsx)",
                            data=excel_data,
                            file_name="cas_de_test.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key="download_tests_excel"
                        )
                    except Exception as e:
                        st.error("Erreur Excel : {}".format(str(e)))

if __name__ == "__main__":
    main()
